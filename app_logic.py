# app_logic.py
import os
import time
import numpy as np
import pandas as pd
from docx import Document
import chromadb
from google import genai
from google.genai import types
from google.api_core import retry

# --- Global API Key (Hardcoded - Use Environment Variables in Production!) ---
GOOGLE_API_KEY = "AIzaSyApn4nBSNFOvf0lB_1b1V2OnNO0qvjWj3Y"
client = genai.Client(api_key=GOOGLE_API_KEY)

# --- Model Configuration ---
model_config = types.GenerateContentConfig(
    temperature=0.75,
    top_p=0.9,
)
search_config= types.GenerateContentConfig(
    tools=[types.Tool(google_search=types.GoogleSearch())],
)

is_retriable = lambda e: (isinstance(e, genai.errors.APIError) and e.code in {429, 503})

def read_docx(endname, folder="Clauses"):
    path = f"//kaggle//input//agreement-clauses-capstone//{folder}//{endname}.docx"
    try:
        doc = Document(path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return paragraphs
    except FileNotFoundError:
        print(f"Error: File not found at {path}")
        return []

def extract_doxc(path):
    try:
        doc = Document(path)
        return doc
    except FileNotFoundError:
        print(f"Error: File not found at {path}")
        return None

def extract_samples(endname, folder="sampleagreemts"):
    dataset_path = f"/kaggle/input/agreement-clauses-capstone/{folder}/{endname}/{endname}"
    docx_files = []
    try:
        for item in os.listdir(dataset_path):
            item_path = os.path.join(dataset_path, item)
            if os.path.isfile(item_path) and item.lower().endswith('.docx'):
                docx_files.append(item_path)
        return docx_files
    except FileNotFoundError:
        print(f"Error: Directory not found at {dataset_path}")
        return []

@retry.Retry(predicate=is_retriable)
def generate_embeddings(cl, etype):
    if etype:
        embedding_task = "retrieval_document"
    else:
        embedding_task = "retrieval_query"
    embed = client.models.embed_content(
        model="models/text-embedding-004",
        contents=cl,
        config=types.EmbedContentConfig(
            task_type=embedding_task
        )
    )
    return [e.values for e in embed.embeddings]

def initialize_chromadb():
    clientdb = chromadb.Client()
    rent = clientdb.get_or_create_collection(name="rent_agreements")
    nda = clientdb.get_or_create_collection(name="nda_agreements")
    employment = clientdb.get_or_create_collection(name="employ_agreements")
    franchise = clientdb.get_or_create_collection(name="franchise_agreements")
    contractor = clientdb.get_or_create_collection(name="contractor_agreements")
    return clientdb, {"rent": rent, "nda": nda, "employment": employment, "franchise": franchise, "contractor": contractor}

def load_agreement_clauses():
    rent_clauses = read_docx("rent")
    nda_clauses = read_docx("nda")
    employment_clauses = read_docx("employment")
    franchise_clauses = read_docx("franchise")
    contractor_clauses = read_docx("contractor")
    return {"rent": rent_clauses, "nda": nda_clauses, "employment": employment_clauses, "franchise": franchise_clauses, "contractor": contractor_clauses}

def store_embeddings_in_chromadb(all_clauses, all_dbs):
    for j, (agreement_type, dataset) in enumerate(all_clauses.items()):
        embeds = []
        ids = []
        documents = []
        db = all_dbs[agreement_type]
        for i, clause in enumerate(dataset):
            vector = generate_embeddings([clause], True)
            time.sleep(0.4) # Added for API call exceeded errors
            embeds.append(vector[0])
            ids.append(f"clause-{agreement_type}-{i}")
            documents.append(clause)
        if embeds:  # Only add if there are embeddings
            db.add(embeddings=embeds, ids=ids, documents=documents)

def strip_type(agr: str):
    agreement_types = ["rent", "nda", "contractor", "employment", "franchise"]
    prompt = f"""Return the type of agreement that the user is referring to in his input {agr}. Respond in one word, all lowercase. Your responses can only be from the set {agreement_types}. Do not use any punctuation. Just respond with the single word."""
    full_prompt = f"""    Prompt: {prompt}        Possible responses: {agreement_types}    Sentence: {agr}        Respond in one word, only with the type. all lowercase. No punctuation    """
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=full_prompt,
        config=model_config
    )
    return response.text.strip().lower()

def pos_neg(response: str):
    prompt = f"""    Classify the sentiment of the following sentence.     Reply with ONLY '1' if the sentence is positive and ONLY '0' if the sentence is negative.    Sentence = {response}    """
    response_heat = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config=model_config
    )
    return bool(int(response_heat.text))

def perform_analysis(atype, impt):
    prompt = f""" You are a legal assistant specialising in determining if the input parameters by the user, defined in {impt} are enough parameters to format an agreement of type {atype}.         Please evaluate if the provided information seems to cover all the generally important aspects for a '{atype}' agreement.    Make sure to evaluate the quality of the input too, if the input seems vague, do consider it as a invalid/bad input.    Your evaluation must be strict and precise. Any vagueness/lack of information must be considered a defect.    However, a extremely subtle nuance must be ignored, for the benefit of the user and for ease of use. Make sure that the input is comprehensive enough to generate a working agreement.    Respond with ONLY these messages, no others.    - "Yes. All essential information seems to be present." if the input appears comprehensive.    - "No, The following essential information seems to be missing or unclear: [list of missing/unclear aspects]" if key details appear to be absent. Be specific about what's lacking (e.g., "names of all parties", "duration of the agreement", "specific details about the confidential information").    - "No, The provided information is too vague or insufficient." if the input is very brief or lacks substantial details.    """
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config=model_config
    )
    return response.text, pos_neg(response.text)

def obtain_information_holes(important_info, extra_info, final_type):
    total_info = important_info + extra_info
    prompt = f"""        The total information given by the user as an input to generate the agreement of type {final_type} are given in {total_info}        Identify any missing or unclear information needed to generate a {final_type} agreement based on the provided user input: {total_info}.         As a comprehensive legal assistant, pinpoint specific details that require clarification or are absent from the input.        Generate your final prompt in a way such that if it is passed into a google search, it gives back the required information.    """
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config=model_config,
    )
    return response.text

def get_data(holes: str, final_type: str):
    prompt= f"""    As a LLM, you have identified a few information deficiencies, outlined in {holes} required to generate a LAW agreement of type {final_type}.    You are supposed to retrieve the relevant information using google search. Make sure to keep it concise and accurate.    """
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
        config=search_config,
    )
    return response.text

def generate_agreement(important_info, extra_info, final_type, relevant_documents, sample_agreements, obtained_info):
    prompt = f"""You are a helpful AI assistant for law agreement generation.The names, dates, locations and information are stored in {important_info} The agreement type is {final_type}.  {relevant_documents} contains most common used clauses in the current agreements, with relevance sorted from highest to lowest, depending on this current use case. Make sure to read through them, understand them,  and use the most relevant documents according to the user's wish as outlined in {extra_info}. Make sure your clauses are end to end, non manipulatable, unable to put loopholes through, and make them concise and readable. While referring to government  officers, refer to them as specifically as possible to avoid confusion. A few example agreement formats are outlined in {sample_agreements}. Structure them similarly and provide a concise output. The english must be clean, non-confusing and clear enough to be understood by a common man, but complex enough to withhold legal intricacies and important points.Some important information that has been obtained through google search is given in {obtained_info}. Make sure to use them in your generation as requested by the user. Format a full {final_type} agreement as provided in the samples and give the final output. Only give the formatted {final_type} agreement in the output. Nothing else."""
    combined_prompt = f"""{prompt} Sample Agreements: {sample_agreements} Relevant Clauses: {relevant_documents}"""
    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=combined_prompt,
        config=model_config
    )
    return response.text

# --- Initialization (moved here for potential direct execution) ---
clientdb, all_dbs = initialize_chromadb()
all_clauses_data = load_agreement_clauses()
store_embeddings_in_chromadb(all_clauses_data, all_dbs)
